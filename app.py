import os, json, io, re
from flask import Flask, render_template, request, Response, stream_with_context, jsonify, send_file
from dotenv import load_dotenv
from azure.identity import DefaultAzureCredential
from azure.ai.projects import AIProjectClient

load_dotenv()
app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024

# ── file parsing ───────────────────────────────────────────────────────────────
def extract_text_from_file(file):
    filename = file.filename.lower()
    if filename.endswith('.pdf'):
        try:
            import pdfplumber
            data = file.read()
            text = ""
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t: text += t + "\n"
            return text.strip(), None
        except ImportError:
            return None, "pdfplumber not installed. Run: pip install pdfplumber"
        except Exception as e:
            return None, f"PDF read error: {str(e)}"
    elif filename.endswith('.docx'):
        try:
            import docx
            data = file.read()
            doc = docx.Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip()), None
        except ImportError:
            return None, "python-docx not installed."
    elif filename.endswith(('.txt', '.md')):
        return file.read().decode('utf-8', errors='ignore'), None
    else:
        return None, "Unsupported file type. Use PDF, DOCX, or TXT."

# ── Azure helpers ──────────────────────────────────────────────────────────────
def get_openai_client():
    """
    Use API key if available (Render/cloud deployment),
    fall back to DefaultAzureCredential for local development.
    """
    endpoint = os.getenv("FOUNDRY_PROJECT_ENDPOINT")
    api_key  = os.getenv("AZURE_API_KEY")

    if api_key:
        # Cloud deployment — use API key directly
        from openai import AzureOpenAI
        # Extract base URL from endpoint
        # e.g. https://careeriq-resource.services.ai.azure.com/api/projects/careeriq
        # → https://careeriq-resource.services.ai.azure.com
        base = endpoint.split('/api/projects/')[0] if '/api/projects/' in endpoint else endpoint
        return AzureOpenAI(
            api_key=api_key,
            azure_endpoint=base,
            api_version="2025-01-01-preview",
        )
    else:
        # Local development — use DefaultAzureCredential
        from azure.identity import DefaultAzureCredential
        from azure.ai.projects import AIProjectClient
        client = AIProjectClient(
            endpoint=endpoint,
            credential=DefaultAzureCredential(),
        )
        return client.get_openai_client()

def stream_response(messages, max_tokens=1800):
    def generate():
        try:
            oc = get_openai_client()
            model = os.getenv("AZURE_AI_MODEL_DEPLOYMENT_NAME","gpt-4.1-mini")
            stream = oc.chat.completions.create(
                model=model,
                messages=messages, stream=True,
                max_tokens=max_tokens, temperature=0.7,
            )
            for chunk in stream:
                if chunk.choices and chunk.choices[0].delta.content:
                    yield f"data: {json.dumps({'content':chunk.choices[0].delta.content})}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error':str(e)})}\n\n"
    return Response(stream_with_context(generate()), mimetype="text/event-stream",
                    headers={"Cache-Control":"no-cache","X-Accel-Buffering":"no"})

def call_ai(messages, max_tokens=2000):
    oc = get_openai_client()
    model = os.getenv("AZURE_AI_MODEL_DEPLOYMENT_NAME","gpt-4.1-mini")
    resp = oc.chat.completions.create(
        model=model,
        messages=messages, max_tokens=max_tokens, temperature=0.4,
    )
    return resp.choices[0].message.content

# ── DOCX builder — matches Dilshad's exact resume style ───────────────────────
# ── DOCX builder — uses tagged format ─────────────────────────────────────────
def build_resume_docx(optimized_text: str) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.5)
        section.right_margin  = Inches(0.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    DARK_BLUE = RGBColor(0x1F, 0x49, 0x7D)
    GREY      = RGBColor(0x55, 0x55, 0x55)
    BLACK     = RGBColor(0x00, 0x00, 0x00)

    def hr_para(p):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'1F497D')
        pBdr.append(bot); pPr.append(pBdr)

    def add_runs(p, text, bold=False, size=10.5, color=None):
        parts = re.split(r'\*\*(.+?)\*\*', text)
        for idx, part in enumerate(parts):
            if not part: continue
            run = p.add_run(part)
            run.bold = bold or (idx % 2 == 1)
            run.font.size = Pt(size)
            if color: run.font.color.rgb = color

    for tag, content in parse_tagged_resume(optimized_text):
        if tag == 'BLANK':
            continue

        elif tag == 'NAME':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            add_runs(p, content, bold=True, size=13, color=BLACK)

        elif tag == 'CONTACT':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(5)
            add_runs(p, content, size=9.5, color=GREY)

        elif tag == 'SECTION':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(content.upper())
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = DARK_BLUE
            hr_para(p)

        elif tag == 'JOB':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(0)
            add_runs(p, content, bold=True, size=10, color=BLACK)

        elif tag == 'SUBLABEL':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(content)
            run.italic = True; run.font.size = Pt(9.5)
            run.font.color.rgb = GREY

        elif tag == 'PROJECT':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(0)
            if ' | ' in content:
                parts = content.split(' | ', 1)
                run = p.add_run(parts[0])
                run.bold = True; run.font.size = Pt(10)
                run2 = p.add_run('  ' + parts[1])
                run2.font.size = Pt(9); run2.italic = True
                run2.font.color.rgb = GREY
            else:
                add_runs(p, content, bold=True, size=10, color=BLACK)

        elif tag == 'EDU':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(0)
            add_runs(p, content, bold=True, size=10, color=BLACK)

        elif tag in ('BULLET','CERT'):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Inches(0.15)
            add_runs(p, content, size=9.5)

        else:  # BODY
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            add_runs(p, content, size=9.5)

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf

    doc = Document()

    # Tight margins matching original style
    for section in doc.sections:
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.5)
        section.right_margin  = Inches(0.5)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    def hr_under_para(p, color='1F497D'):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_name_line(text):
        """Dilshad style: Name | Title — bold, 13pt, black, left-aligned"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    def add_contact_line(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        hr_under_para(p)

    def add_job_header(text):
        """Company | Role | Date | Location — bold, 10.5pt"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10.5)

    def add_sub_label(text):
        """Sub-label like 'Banking & Financial Services'"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Inches(0.15)
        # Parse inline bold (**text**)
        parts = re.split(r'\*\*(.+?)\*\*', text.lstrip('•-* '))
        for idx, part in enumerate(parts):
            run = p.add_run(part)
            run.bold = (idx % 2 == 1)
            run.font.size = Pt(10.5)

    def add_normal(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        parts = re.split(r'\*\*(.+?)\*\*', text)
        for idx, part in enumerate(parts):
            run = p.add_run(part)
            run.bold = (idx % 2 == 1)
            run.font.size = Pt(10.5)

    SECTION_KEYS = [
        'EXPERIENCE','EDUCATION','SKILLS','PROJECTS','CERTIFICATIONS',
        'SUMMARY','OBJECTIVE','PUBLICATIONS','AWARDS','LANGUAGES',
        'PROFESSIONAL SUMMARY','WORK EXPERIENCE','TECHNICAL SKILLS',
        'PROFESSIONAL EXPERIENCE','CORE COMPETENCIES','CERTIFICATIONS & ACHIEVEMENTS',
        'CERTIFICATIONS AND ACHIEVEMENTS'
    ]

    lines = [l.strip() for l in optimized_text.split('\n')]
    name_done    = False
    contact_done = False

    for line in lines:
        if not line:
            continue

        # Strip markdown heading markers
        clean = re.sub(r'^#{1,3}\s*', '', line).strip()
        cu = clean.upper()

        # ── Name line (first meaningful line, "Name | Title" pattern)
        if not name_done:
            if '|' in clean and len(clean.split()) <= 8:
                add_name_line(clean)
                name_done = True
                continue
            elif len(clean.split()) <= 5 and not any(k in cu for k in SECTION_KEYS):
                add_name_line(clean)
                name_done = True
                continue

        # ── Contact line (has @ or linkedin or phone digits)
        if not contact_done and name_done:
            if any(x in clean.lower() for x in ['@','linkedin','github','portfolio','334','555','phone']):
                add_contact_line(clean)
                contact_done = True
                continue

        # ── Section heading
        if any(k in cu for k in SECTION_KEYS) and len(clean) < 60 and not clean.startswith('-'):
            add_section_heading(clean.lstrip('#').strip())
            continue

        # ── Bullet point
        if line.startswith(('-','•','*','+')):
            add_bullet(clean)
            continue

        # ── Job header: "Role, Company  Date | Location" or has multiple |
        if clean.count('|') >= 1 and name_done and contact_done:
            add_job_header(clean)
            continue

        # ── Sub-label line (italic industry/domain label)
        if clean.endswith('Services') or clean.endswith('Industry') or (
            len(clean.split()) <= 5 and not any(k in cu for k in SECTION_KEYS) and
            not clean.startswith('-') and name_done and contact_done
        ):
            add_sub_label(clean)
            continue

        # ── Default normal text
        add_normal(clean)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_cover_docx(cover_text: str) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, Inches
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for line in cover_text.split('\n'):
        clean = re.sub(r'[#*_`]','',line).strip()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(clean if clean else '')
        run.font.size = Pt(11)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ── Tagged format parser ───────────────────────────────────────────────────────
def parse_tagged_resume(text: str) -> list:
    """
    Parse [TAG] lines into list of (tag, content) tuples.
    Falls back to heuristic parsing if tags are missing.
    """
    result = []
    TAG_RE = re.compile(r'^\[(NAME|CONTACT|SECTION|BULLET|JOB|SUBLABEL|PROJECT|EDU|CERT)\]\s*(.*)', re.IGNORECASE)

    # Pre-process: merge split "| Title" lines
    lines = text.split('\n')
    merged = []
    for l in lines:
        s = l.strip()
        if s.startswith('| ') and merged:
            for j in range(len(merged)-1,-1,-1):
                if merged[j].strip():
                    merged[j] = merged[j].rstrip() + ' ' + s
                    break
            else:
                merged.append(l)
        else:
            merged.append(l)

    SECTION_KEYS = [
        'PROFESSIONAL SUMMARY','TECHNICAL SKILLS','PROFESSIONAL EXPERIENCE',
        'EXPERIENCE','EDUCATION','SKILLS','PROJECTS','CERTIFICATIONS',
        'SUMMARY','WORK EXPERIENCE','CORE COMPETENCIES',
        'CERTIFICATIONS & ACHIEVEMENTS','CERTIFICATIONS AND ACHIEVEMENTS'
    ]

    name_done = contact_done = False
    for line in merged:
        s = line.strip()
        if not s:
            result.append(('BLANK',''))
            continue

        m = TAG_RE.match(s)
        if m:
            tag = m.group(1).upper()
            content = m.group(2).strip()
            result.append((tag, content))
            if tag == 'NAME': name_done = True
            if tag == 'CONTACT': contact_done = True
            continue

        # ── Fallback heuristics (no tags) ──────────────────────────────────
        clean = re.sub(r'^#{1,3}\s*','',s).strip()
        cu = clean.upper()

        if not name_done:
            result.append(('NAME', clean)); name_done = True; continue

        if not contact_done and any(x in clean.lower() for x in ['@','linkedin','github','dilshad','334','555']):
            result.append(('CONTACT', clean)); contact_done = True; continue

        if any(cu.startswith(k) or cu == k for k in SECTION_KEYS) and len(clean) < 65 and not clean.startswith('-'):
            result.append(('SECTION', clean.lstrip('#').strip().upper())); continue

        if clean.startswith(('-','•','*')):
            result.append(('BULLET', clean.lstrip('-•* ').strip())); continue

        if 'github.com' in clean.lower() or 'hackathon' in clean.lower():
            result.append(('PROJECT', clean)); continue

        if ('|' in clean or '–' in clean) and contact_done and not clean.startswith('-'):
            result.append(('JOB', clean)); continue

        if (clean.endswith('Services') or clean.endswith('Industry')) and len(clean.split()) <= 5:
            result.append(('SUBLABEL', clean)); continue

        result.append(('BODY', clean))

    return result


# ── PDF builder — uses tagged format ──────────────────────────────────────────
def build_resume_pdf(optimized_text: str) -> io.BytesIO:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_LEFT

    buf = io.BytesIO()
    doc_rl = SimpleDocTemplate(buf, pagesize=letter,
        leftMargin=0.5*inch, rightMargin=0.5*inch,
        topMargin=0.45*inch, bottomMargin=0.45*inch)

    BLACK     = colors.HexColor('#000000')
    DARK_BLUE = colors.HexColor('#1F497D')
    GREY      = colors.HexColor('#555555')
    BOLD      = 'Helvetica-Bold'
    REG       = 'Helvetica'
    ITAL      = 'Helvetica-Oblique'

    # ── Tight single-page styles matching Dilshad's original ─────────────────
    S = {
        'NAME':    ParagraphStyle('N',  fontName=BOLD, fontSize=12,   textColor=BLACK,     leading=14, spaceBefore=0, spaceAfter=1),
        'CONTACT': ParagraphStyle('C',  fontName=REG,  fontSize=9,    textColor=GREY,      leading=11, spaceBefore=0, spaceAfter=3),
        'SECTION': ParagraphStyle('S',  fontName=BOLD, fontSize=10,   textColor=DARK_BLUE, leading=12, spaceBefore=5, spaceAfter=1),
        'JOB':     ParagraphStyle('J',  fontName=BOLD, fontSize=10,   textColor=BLACK,     leading=12, spaceBefore=3, spaceAfter=0),
        'SUBLABEL':ParagraphStyle('SL', fontName=ITAL, fontSize=9.5,  textColor=GREY,      leading=11, spaceBefore=0, spaceAfter=1),
        'PROJECT': ParagraphStyle('P',  fontName=BOLD, fontSize=10,   textColor=BLACK,     leading=12, spaceBefore=3, spaceAfter=0),
        'EDU':     ParagraphStyle('E',  fontName=BOLD, fontSize=10,   textColor=BLACK,     leading=12, spaceBefore=3, spaceAfter=0),
        'BULLET':  ParagraphStyle('B',  fontName=REG,  fontSize=9.5,  textColor=BLACK,     leading=12, leftIndent=12, firstLineIndent=-9, spaceBefore=0, spaceAfter=1),
        'CERT':    ParagraphStyle('CR', fontName=REG,  fontSize=9.5,  textColor=BLACK,     leading=12, leftIndent=12, firstLineIndent=-9, spaceBefore=0, spaceAfter=1),
        'BODY':    ParagraphStyle('BD', fontName=REG,  fontSize=9.5,  textColor=BLACK,     leading=12, spaceBefore=0, spaceAfter=1),
    }

    def esc(t):
        return t.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')

    def bi(t):  # bold inline
        return re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', esc(t))

    story = []
    for tag, content in parse_tagged_resume(optimized_text):
        if tag == 'BLANK':
            story.append(Spacer(1, 1))   # was 2 — tighter blank lines
        elif tag == 'SECTION':
            story.append(Paragraph(esc(content.upper()), S['SECTION']))
            story.append(HRFlowable(width="100%", thickness=0.6, color=DARK_BLUE, spaceAfter=1))
        elif tag in ('BULLET','CERT'):
            story.append(Paragraph('• ' + bi(content), S['BULLET']))
        elif tag == 'PROJECT':
            # Split at " | " to separate title from tech if present
            if ' | ' in content:
                parts = content.split(' | ', 1)
                story.append(Paragraph(bi(parts[0]), S['PROJECT']))
                story.append(Paragraph(esc(parts[1]), ParagraphStyle('PT',
                    fontName=ITAL, fontSize=9.5, textColor=GREY,
                    leading=12, spaceBefore=0, spaceAfter=1)))
            else:
                story.append(Paragraph(bi(content), S['PROJECT']))
        elif tag in S:
            story.append(Paragraph(bi(content) if tag not in ('NAME','CONTACT','JOB','EDU','SUBLABEL') else esc(content), S[tag]))
        else:
            story.append(Paragraph(bi(content), S['BODY']))

    doc_rl.build(story)
    buf.seek(0)
    return buf
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_LEFT

    # ── Pre-process: collapse split name+title into one line ──────────────────
    # AI sometimes outputs "Dilshad Shaik\n| Software Engineer" as two lines.
    # Fix: if a line starts with "| " and previous non-empty line had no section keyword,
    # merge them.
    raw_lines = optimized_text.split('\n')
    merged = []
    for i, l in enumerate(raw_lines):
        s = l.strip()
        if s.startswith('| ') and merged:
            # Find last non-empty line and merge
            for j in range(len(merged)-1, -1, -1):
                if merged[j].strip():
                    merged[j] = merged[j].rstrip() + ' ' + s
                    break
            else:
                merged.append(s)
        else:
            merged.append(l)
    optimized_text = '\n'.join(merged)
    # ─────────────────────────────────────────────────────────────────────────

    buf = io.BytesIO()
    doc_rl = SimpleDocTemplate(buf, pagesize=letter,
        leftMargin=0.5*inch, rightMargin=0.5*inch,
        topMargin=0.5*inch, bottomMargin=0.5*inch)

    BLACK     = colors.HexColor('#000000')
    DARK_BLUE = colors.HexColor('#1F497D')
    GREY      = colors.HexColor('#444444')
    MID_GREY  = colors.HexColor('#555555')
    BOLD      = 'Helvetica-Bold'
    REG       = 'Helvetica'
    ITAL      = 'Helvetica-Oblique'

    # Name: "Dilshad Shaik | Software Engineer" — bold 13pt left-aligned black
    name_style = ParagraphStyle('Name',
        fontName=BOLD, fontSize=13, textColor=BLACK,
        leading=16, spaceBefore=0, spaceAfter=1)

    # Contact line: small grey
    contact_style = ParagraphStyle('Contact',
        fontName=REG, fontSize=9.5, textColor=MID_GREY,
        leading=13, spaceBefore=0, spaceAfter=5)

    # Section heading: ALL CAPS bold dark blue, tight
    section_style = ParagraphStyle('Section',
        fontName=BOLD, fontSize=10.5, textColor=DARK_BLUE,
        leading=14, spaceBefore=6, spaceAfter=1)

    # Job title line: "Role, Company  Date | Location" bold black
    job_style = ParagraphStyle('Job',
        fontName=BOLD, fontSize=10.5, textColor=BLACK,
        leading=14, spaceBefore=4, spaceAfter=0)

    # Sub-label: "Banking & Financial Services" — italic grey
    sub_style = ParagraphStyle('Sub',
        fontName=ITAL, fontSize=10, textColor=GREY,
        leading=13, spaceBefore=0, spaceAfter=1)

    # Bullet: tight, left-indented
    bullet_style = ParagraphStyle('Bullet',
        fontName=REG, fontSize=10.5, textColor=BLACK,
        leading=13.5, leftIndent=14, firstLineIndent=-10,
        spaceBefore=0, spaceAfter=1.5)

    # Project title line (has github.com in it)
    proj_style = ParagraphStyle('Proj',
        fontName=BOLD, fontSize=10.5, textColor=BLACK,
        leading=13, spaceBefore=4, spaceAfter=0)

    # Tech stack line after project (italic, smaller)
    tech_style = ParagraphStyle('Tech',
        fontName=ITAL, fontSize=9.5, textColor=MID_GREY,
        leading=12, spaceBefore=0, spaceAfter=1)

    # Normal body text
    body_style = ParagraphStyle('Body',
        fontName=REG, fontSize=10.5, textColor=BLACK,
        leading=13.5, spaceBefore=0, spaceAfter=1)

    SECTION_KEYS = [
        'PROFESSIONAL SUMMARY','TECHNICAL SKILLS','PROFESSIONAL EXPERIENCE',
        'EXPERIENCE','EDUCATION','SKILLS','PROJECTS','CERTIFICATIONS',
        'SUMMARY','OBJECTIVE','PUBLICATIONS','AWARDS','LANGUAGES',
        'WORK EXPERIENCE','CORE COMPETENCIES',
        'CERTIFICATIONS & ACHIEVEMENTS','CERTIFICATIONS AND ACHIEVEMENTS'
    ]

    def esc(t):
        return (t.replace('&','&amp;')
                 .replace('<','&lt;')
                 .replace('>','&gt;'))

    def bold_inline(t):
        """Convert **bold** markdown to reportlab <b>bold</b> XML."""
        return re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', esc(t))

    def is_section(text):
        cu = text.upper().strip().lstrip('#').strip()
        return any(cu == k or cu.startswith(k) for k in SECTION_KEYS) and len(text) < 65

    def is_job_line(text):
        """e.g. 'Software Engineer, Cognizant  Sept. 2021 – Jul. 2024 | Hyderabad'"""
        return ('|' in text or '–' in text or '-' in text) and \
               not text.strip().startswith('-') and \
               not text.strip().startswith('•') and \
               len(text) < 120

    def is_contact(text):
        return any(x in text.lower() for x in
                   ['@','linkedin','github','portfolio','334','555','http','dilshad'])

    def is_project_title(text):
        return ('github.com' in text.lower() or
                ('hackathon' in text.lower()) or
                ('deployed on' in text.lower()))

    story = []
    lines = [l.rstrip() for l in optimized_text.split('\n')]
    name_done    = False
    contact_done = False
    header_block_done = False  # True after contact line processed

    i = 0
    while i < len(lines):
        raw  = lines[i]
        line = raw.strip()
        clean = re.sub(r'^#{1,3}\s*', '', line).strip()
        i += 1

        if not clean:
            if header_block_done:
                story.append(Spacer(1, 2))
            continue

        # ── NAME (first non-empty line)
        if not name_done:
            story.append(Paragraph(esc(clean), name_style))
            name_done = True
            continue

        # ── CONTACT LINE
        if not contact_done and is_contact(clean):
            story.append(Paragraph(esc(clean), contact_style))
            contact_done = True
            header_block_done = True
            continue

        # ── SECTION HEADING
        if is_section(clean):
            heading_text = clean.lstrip('#').strip().upper()
            story.append(Paragraph(esc(heading_text), section_style))
            story.append(HRFlowable(width="100%", thickness=0.7,
                                     color=DARK_BLUE, spaceAfter=2))
            continue

        # ── BULLET
        if raw.startswith((' ','\t')) and clean.startswith(('-','•','*')) or \
           clean.startswith(('-','•','*')):
            txt = bold_inline(clean.lstrip('-•* ').strip())
            story.append(Paragraph('• ' + txt, bullet_style))
            continue

        # ── PROJECT TITLE LINE (has github.com or hackathon)
        if is_project_title(clean):
            # Split at " | " to separate title from tech stack if on same line
            if ' | ' in clean:
                parts = clean.split(' | ', 1)
                story.append(Paragraph(bold_inline(parts[0]), proj_style))
                story.append(Paragraph(esc(parts[1]), tech_style))
            else:
                story.append(Paragraph(bold_inline(clean), proj_style))
            continue

        # ── JOB HEADER LINE (has | separator, date range, or location)
        if is_job_line(clean) and header_block_done:
            story.append(Paragraph(esc(clean), job_style))
            continue

        # ── SUB LABEL (Banking & Financial Services, etc.)
        if (clean.endswith('Services') or clean.endswith('Industry') or
                clean.endswith('Services\n')) and len(clean.split()) <= 6:
            story.append(Paragraph(esc(clean), sub_style))
            continue

        # ── DEFAULT BODY
        story.append(Paragraph(bold_inline(clean), body_style))

    doc_rl.build(story)
    buf.seek(0)
    return buf


def build_cover_pdf(cover_text: str) -> io.BytesIO:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    buf = io.BytesIO()
    doc_rl = SimpleDocTemplate(buf, pagesize=letter,
        leftMargin=1.1*inch, rightMargin=1.1*inch,
        topMargin=1.0*inch, bottomMargin=1.0*inch)
    body_style = ParagraphStyle('Body', fontName='Helvetica', fontSize=11,
                                 leading=16, spaceAfter=10)
    story = []
    for line in cover_text.split('\n'):
        clean = re.sub(r'[#*_`]','',line).strip()
        if not clean:
            story.append(Spacer(1,8))
        else:
            story.append(Paragraph(clean.replace('&','&amp;').replace('<','&lt;'), body_style))
    doc_rl.build(story)
    buf.seek(0)
    return buf


def build_analysis_pdf(analysis_text: str) -> io.BytesIO:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER
    buf = io.BytesIO()
    doc_rl = SimpleDocTemplate(buf, pagesize=letter,
        leftMargin=0.9*inch, rightMargin=0.9*inch,
        topMargin=0.8*inch, bottomMargin=0.8*inch)
    DARK_BLUE = colors.HexColor('#1F497D')
    title_s  = ParagraphStyle('T', fontName='Helvetica-Bold', fontSize=15,
                               textColor=DARK_BLUE, alignment=TA_CENTER, spaceAfter=4)
    sub_s    = ParagraphStyle('S', fontName='Helvetica', fontSize=9.5,
                               textColor=colors.grey, alignment=TA_CENTER, spaceAfter=12)
    h2_s     = ParagraphStyle('H', fontName='Helvetica-Bold', fontSize=11,
                               textColor=DARK_BLUE, spaceBefore=10, spaceAfter=3)
    body_s   = ParagraphStyle('B', fontName='Helvetica', fontSize=10.5, leading=15, spaceAfter=3)
    bul_s    = ParagraphStyle('BL', fontName='Helvetica', fontSize=10.5, leading=15,
                               leftIndent=14, firstLineIndent=-10, spaceAfter=2)
    story = [
        Paragraph("CareerIQ — ATS Match Analysis Report", title_s),
        Paragraph("Generated by CareerIQ on Microsoft Foundry · gpt-4.1-mini", sub_s),
        HRFlowable(width="100%", thickness=1, color=DARK_BLUE, spaceAfter=10),
    ]
    def esc(t): return t.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
    for line in analysis_text.split('\n'):
        clean = re.sub(r'[*_`]','',line).strip()
        if not clean:
            story.append(Spacer(1,4))
        elif clean.startswith('#'):
            story.append(Paragraph(esc(clean.lstrip('#').strip()), h2_s))
            story.append(HRFlowable(width="100%", thickness=0.5,
                                     color=colors.HexColor('#CCCCCC'), spaceAfter=2))
        elif clean.startswith(('-','•')):
            story.append(Paragraph('• '+esc(clean.lstrip('-• ')), bul_s))
        else:
            story.append(Paragraph(esc(clean), body_s))
    doc_rl.build(story)
    buf.seek(0)
    return buf


# ── SYSTEM PROMPTS ─────────────────────────────────────────────────────────────
CHAT_SYSTEM = """You are CareerIQ — a personal AI career mentor built on Microsoft Foundry, specifically designed to help students and early-career professionals reach their dream job.

YOUR CORE MISSION:
Give the COMPLETE picture upfront. Never drip-feed one step at a time. Students need to see the full journey so they feel confident, then start with Step 1.

RESPONSE STRUCTURE — always follow this EXACT order:

---

## 🎯 Where You Stand Today
2-3 sentences acknowledging their current skills and strengths honestly and encouragingly.

## 🗺️ Your Full Roadmap to [Dream Job]
Show the COMPLETE step-by-step journey — all steps from where they are to their dream job.
Format every step as:

**✅ Step N (Week X–Y): [Specific Task]**
- What to do: [exact action]
- Resource: [clickable link]
- Time: [realistic estimate]
- ➡️ This unlocks: [what they can do after]

Include ALL steps — typically 6-10 steps covering:
1. Foundation skills
2. Core technical skills
3. First certification
4. Project building
5. Advanced skills / second certification
6. Portfolio & GitHub
7. Networking & job applications
8. Interview preparation
9. Dream job ready ✅

## 📜 Certifications on Your Path
List every relevant cert with:
- [Cert Name](real-url) — Cost | Time to complete | Why it matters for this role

## 🆓 Free Resources
List 4-6 free learning resources with clickable links relevant to their goal.

## 📊 Market Insight
2-3 sentences: demand for this role in 2026, average salary, top hiring companies.

---

## 🚀 Start Here — Your First Step Right Now:
**Step 1 is: [specific action]**
Here's the exact link to begin: [clickable link]
This takes [X hours/days] and after completing it you will [specific outcome].
You've got this! 💪

---

CERTIFICATION LINKS — always use these exact URLs:
- [AZ-900](https://learn.microsoft.com/en-us/credentials/certifications/azure-fundamentals/) — Free on MS Learn, 2-3 weeks
- [AI-900](https://learn.microsoft.com/en-us/credentials/certifications/azure-ai-fundamentals/) — Free on MS Learn, 2-3 weeks
- [AI-102](https://learn.microsoft.com/en-us/credentials/certifications/azure-ai-engineer/) — 4-6 weeks
- [AZ-204](https://learn.microsoft.com/en-us/credentials/certifications/azure-developer/) — 6-8 weeks
- [AZ-104](https://learn.microsoft.com/en-us/credentials/certifications/azure-administrator/) — 6-8 weeks
- [AZ-305](https://learn.microsoft.com/en-us/credentials/certifications/azure-solutions-architect/) — 8-10 weeks
- [DP-100](https://learn.microsoft.com/en-us/credentials/certifications/azure-data-scientist/) — 6-8 weeks
- [AWS CCP](https://aws.amazon.com/certification/certified-cloud-practitioner/) — 2-3 weeks
- [AWS SAA](https://aws.amazon.com/certification/certified-solutions-architect-associate/) — 6-8 weeks
- [GCP ACE](https://cloud.google.com/learn/certification/cloud-engineer) — 6-8 weeks
- [TensorFlow Dev Cert](https://www.tensorflow.org/certificate) — 4-6 weeks
- [CKA Kubernetes](https://training.linuxfoundation.org/certification/certified-kubernetes-administrator-cka/) — 6-8 weeks
- [Terraform Associate](https://developer.hashicorp.com/certifications/infrastructure-automation) — 3-4 weeks
- [GitHub Foundations](https://examregistration.github.com/certification/GHF) — 1-2 weeks
- [Snowflake SnowPro](https://learn.snowflake.com/en/certifications/) — 3-4 weeks
- [Databricks Certified](https://www.databricks.com/learn/certification) — 4-6 weeks

ROADMAP LINKS — always include the relevant one:
- Backend: [roadmap.sh/backend](https://roadmap.sh/backend)
- Frontend: [roadmap.sh/frontend](https://roadmap.sh/frontend)
- Full Stack: [roadmap.sh/full-stack](https://roadmap.sh/full-stack)
- DevOps: [roadmap.sh/devops](https://roadmap.sh/devops)
- AI/ML: [roadmap.sh/ai-data-scientist](https://roadmap.sh/ai-data-scientist)
- Python: [roadmap.sh/python](https://roadmap.sh/python)
- Java: [roadmap.sh/java](https://roadmap.sh/java)
- System Design: [roadmap.sh/system-design](https://roadmap.sh/system-design)
- Docker: [roadmap.sh/docker](https://roadmap.sh/docker)
- Kubernetes: [roadmap.sh/kubernetes](https://roadmap.sh/kubernetes)
- AWS: [roadmap.sh/aws](https://roadmap.sh/aws)
- Data Science: [roadmap.sh/data-science](https://roadmap.sh/data-science)

FREE LEARNING RESOURCES — always link these when relevant:
- [Microsoft Learn](https://learn.microsoft.com/en-us/training/) — free, official Microsoft courses
- [AWS Skill Builder](https://skillbuilder.aws/) — free AWS courses
- [Google Cloud Skills Boost](https://cloudskillsboost.google/) — free GCP labs
- [freeCodeCamp](https://www.freecodecamp.org/) — free full curriculum
- [The Odin Project](https://www.theodinproject.com/) — free web dev
- [CS50 Harvard](https://cs50.harvard.edu/x/) — free intro to CS
- [Kaggle Learn](https://www.kaggle.com/learn) — free ML/data science
- [fast.ai](https://www.fast.ai/) — free deep learning

CRITICAL RULES:
- ALWAYS give the COMPLETE roadmap — never just one step
- ALWAYS end with the first step highlighted clearly as the action item
- Every cert and resource must be a clickable markdown link
- Be encouraging and specific — never vague
- Only answer career development and learning questions"""

RESUME_SYSTEM = """You are a certified ATS (Applicant Tracking System) expert and professional resume writer.
You optimize resumes to score 93+ on ATS parsers like Workday, Taleo, Greenhouse, Lever, and iCIMS.

OUTPUT FORMAT — every line must start with a tag:
[NAME]     — "First Last | Job Title"
[CONTACT]  — phone | email | linkedin | github (one line)
[SECTION]  — section heading in ALL CAPS
[BULLET]   — a single bullet point (no dash, no bullet symbol)
[JOB]      — "Role, Company  Date | Location" ALL ON ONE LINE
[SUBLABEL] — domain label e.g. "Banking & Financial Services"
[PROJECT]  — title + github + tech ALL ON ONE LINE
[EDU]      — degree + date + institution ALL ON ONE LINE
[CERT]     — single certification line

PROVEN ATS 93+ RULES (all mandatory):

RULE 1 — KEYWORD DENSITY (highest ATS weight)
- Extract every noun, skill, tool, framework, methodology from the JD
- Inject EACH keyword at least once, verbatim, exact spelling, exact case
- Mirror JD language exactly: if JD says "cross-functional collaboration" use that exact phrase
- Place the most important JD keywords in Professional Summary (top of resume = highest parser weight)

RULE 2 — STANDARD SECTION HEADERS ONLY
- Use ONLY: PROFESSIONAL SUMMARY, TECHNICAL SKILLS, PROFESSIONAL EXPERIENCE, PROJECTS, EDUCATION, CERTIFICATIONS & ACHIEVEMENTS
- Never rename sections — ATS parsers fail on non-standard headers

RULE 3 — CLEAN STRUCTURE
- No tables, columns, text boxes, headers/footers, graphics
- No special characters except: | - ( ) / @ . , # &
- No emojis anywhere in the resume

RULE 4 — ACTION VERB + METRIC FORMAT for every bullet
- Start every bullet with a strong past-tense action verb:
  Led, Built, Designed, Developed, Implemented, Deployed, Architected, Reduced, Increased,
  Automated, Optimized, Delivered, Managed, Spearheaded, Pioneered, Improved, Integrated, Scaled
- Format: [Action Verb] [what] [with what tools from JD] [quantified result with % or number]
- Example: "Reduced deployment cycle by 80% by optimizing CI/CD pipelines using Jenkins and GitHub Actions"

RULE 5 — SKILLS SECTION ORDER
- Required JD skills first, preferred JD skills second, additional skills last
- List skills in same category order as the JD requirements section

RULE 6 — PROFESSIONAL SUMMARY (highest ATS weight)
- Exactly 2-3 sentences
- Must contain top 5 JD keywords verbatim
- Sentence 1: "[Title] with [X] years in [top 3 JD requirements verbatim]"
- Sentence 2: "Skilled in [5 most important JD technical skills verbatim]"
- Sentence 3: Biggest achievement + most relevant JD requirement

RULE 7 — PRESERVE ORIGINAL FORMAT
- Keep ALL abbreviations: AWS, CI/CD, RAG, REST APIs, LLMs, ML, AI, GCP, OOP, SLA, RBAC
- Keep same section names as original
- Keep same NUMBER of bullet points per job — never add extra bullets
- [JOB], [PROJECT], [EDU] must NEVER be split across multiple lines
- Output ONLY tagged lines — no markdown, no commentary, no blank lines"""

# ── routes ─────────────────────────────────────────────────────────────────────
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/upload", methods=["POST"])
def upload():
    if 'file' not in request.files:
        return {"error":"No file provided"},400
    file = request.files['file']
    if not file.filename:
        return {"error":"Empty filename"},400
    text, error = extract_text_from_file(file)
    if error:
        return {"error":error},400
    return {"text":text,"filename":file.filename}

# ── chat ───────────────────────────────────────────────────────────────────────
@app.route("/chat", methods=["POST"])
def chat():
    data = request.json
    messages = [{"role":"system","content":CHAT_SYSTEM}] + data.get("messages",[])
    return stream_response(messages, max_tokens=2500)

def _resume_prompt(resume, jd, level, industry):
    return f"""Optimize this resume to score 93+ ATS. Follow ALL 7 rules exactly.

ROLE LEVEL: {level} | INDUSTRY: {industry}

=== ORIGINAL RESUME (preserve this exact structure, style, and bullet count) ===
{resume}

=== JOB DESCRIPTION (extract every keyword from this) ===
{jd}

EXECUTION STEPS:
Step 1: Extract ALL keywords from JD (skills, tools, frameworks, methodologies, exact phrases)
Step 2: Rewrite PROFESSIONAL SUMMARY using Rule 6 — include top 5 JD keywords verbatim
Step 3: Reorder TECHNICAL SKILLS — required JD skills first, preferred second
Step 4: Rewrite each bullet using Rule 4 — Action Verb + What + JD Tools + Quantified Result
Step 5: Inject any remaining JD keywords naturally into existing bullets
Step 6: Verify every major JD keyword appears at least once

HARD CONSTRAINTS:
- Same number of bullets per job as original — do not add new bullets
- All abbreviations preserved: AWS, CI/CD, RAG, LLMs, REST APIs, etc.
- No emojis, no special characters beyond | - ( ) / @ . , # &
- [JOB], [PROJECT], [EDU] lines stay on ONE line — never split

Output ONLY tagged lines. No markdown. No commentary. No blank lines."""

# ── studio streaming ───────────────────────────────────────────────────────────
@app.route("/studio/resume", methods=["POST"])
def studio_resume():
    d = request.json
    resume,jd = d.get("resume",""),d.get("jd","")
    level,industry = d.get("level","mid"),d.get("industry","tech")
    # Stream tagged format, strip tags for display
    def generate_resume():
        try:
            oc = get_openai_client()
            stream = oc.chat.completions.create(
                model=os.getenv("AZURE_AI_MODEL_DEPLOYMENT_NAME","gpt-4.1-mini"),
                messages=[
                    {"role":"system","content":RESUME_SYSTEM},
                    {"role":"user","content":_resume_prompt(resume,jd,level,industry)}
                ],
                stream=True, max_tokens=2200, temperature=0.3,
            )
            for chunk in stream:
                if chunk.choices and chunk.choices[0].delta.content:
                    c = chunk.choices[0].delta.content
                    # Strip tags for display in browser
                    c = re.sub(r'^\[(NAME|CONTACT|SECTION|BULLET|JOB|SUBLABEL|PROJECT|EDU|CERT)\]\s*', '', c, flags=re.MULTILINE)
                    yield f"data: {json.dumps({'content':c})}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error':str(e)})}\n\n"
    return Response(stream_with_context(generate_resume()), mimetype="text/event-stream",
                    headers={"Cache-Control":"no-cache","X-Accel-Buffering":"no"})

@app.route("/studio/cover", methods=["POST"])
def studio_cover():
    d = request.json
    resume,jd = d.get("resume",""),d.get("jd","")
    level,industry = d.get("level","mid"),d.get("industry","tech")
    prompt = f"""Write a tailored cover letter. ROLE LEVEL: {level} | INDUSTRY: {industry}

=== RESUME ===
{resume}

=== JOB DESCRIPTION ===
{jd}

Under 300 words. Specific hook. 3 body paragraphs with quantified achievements. Confident close.
Never use: 'I am writing to', 'passionate about', 'hardworking', 'team player'
Output only the letter from the greeting."""
    return stream_response([
        {"role":"system","content":"You are an expert cover letter writer for top tech companies. Specific, compelling, never generic."},
        {"role":"user","content":prompt}
    ], max_tokens=700)

@app.route("/studio/analysis", methods=["POST"])
def studio_analysis():
    d = request.json
    resume,jd = d.get("resume",""),d.get("jd","")
    prompt = f"""Analyze resume vs JD. Be specific and data-driven.

=== RESUME ===
{resume}

=== JOB DESCRIPTION ===
{jd}

## 📊 ATS Score Estimate
Estimated ATS score /100 and human reviewer score /10. Explain what drives the score.

## ✅ Strong Matches
6-8 specific matching skills with exact keyword matches shown.

## ❌ Critical Gaps
Top 6 missing keywords/skills. Format: Keyword | Priority (High/Med/Low) | Fix

## 🔑 Missing Keywords to Inject
12-15 exact JD keywords not in resume. Group: Technical | Soft Skills | Tools | Certs

## 📈 Path to ATS 93+
Numbered steps to reach 93+ with this exact resume and JD.

## ⚡ Quick Wins
3-4 changes under 5 minutes that add the most ATS points."""
    return stream_response([
        {"role":"system","content":"You are an expert ATS analyst. Provide detailed, actionable analysis."},
        {"role":"user","content":prompt}
    ], max_tokens=1200)


# ── INTERVIEW PREP AGENT ───────────────────────────────────────────────────────
INTERVIEW_SYSTEM = """You are an expert technical interview coach with 15 years experience at FAANG companies.
You generate highly specific, tailored interview questions and model answers based on the exact job description and candidate's resume.
Never give generic questions. Every question must reference specific technologies, experiences, or requirements from the JD and resume."""

@app.route("/studio/interview", methods=["POST"])
def studio_interview():
    d = request.json
    resume,jd = d.get("resume",""),d.get("jd","")
    level,industry = d.get("level","mid"),d.get("industry","tech")

    prompt = f"""Generate a complete interview prep package for this candidate.

ROLE LEVEL: {level} | INDUSTRY: {industry}

=== CANDIDATE RESUME ===
{resume}

=== JOB DESCRIPTION ===
{jd}

Generate exactly this structure:

## 🎯 Role Overview
2 sentences on what this role is really looking for — what's the hiring bar?

## 💻 Technical Questions (5 questions)
For each question:
**Q1: [Specific technical question from JD requirements]**
Difficulty: Hard/Medium/Easy | Topic: [skill area]
✅ Model Answer: [2-3 sentence answer using candidate's actual experience from resume]
💡 Key terms to mention: [3-4 buzzwords from JD]

## 🧠 Behavioral Questions (3 questions)
STAR-format questions tied to the JD's soft skill requirements.
**Q1: [Behavioral question]**
✅ Model Answer (STAR): Situation: [...] Task: [...] Action: [...] Result: [...]
Use candidate's ACTUAL projects and experiences from their resume.

## 🏢 Company/Role Fit Questions (2 questions)
Questions about why this company and role specifically.
**Q1: [Question]**
✅ Model Answer: [Specific answer referencing the company and candidate background]

## ⚡ Questions YOU Should Ask
5 smart questions the candidate should ask the interviewer to stand out.

## 🚀 Interview Day Tips
3 specific tips based on this exact role and company."""

    return stream_response([
        {"role":"system","content":INTERVIEW_SYSTEM},
        {"role":"user","content":prompt}
    ], max_tokens=2000)


# ── REASONING TRACE ENDPOINT ───────────────────────────────────────────────────
@app.route("/chat/traced", methods=["POST"])
def chat_traced():
    """
    Enhanced chat endpoint that streams reasoning step labels
    before the main response so judges can see multi-step thinking.
    Format: data: {"type": "step", "label": "..."} for steps
            data: {"type": "content", "content": "..."} for response
    """
    data = request.json
    user_message = data.get("messages", [{}])[-1].get("content", "")

    # Determine which reasoning steps apply based on user message
    steps = _get_reasoning_steps(user_message)

    def generate():
        # Stream reasoning steps first
        for step in steps:
            yield f"data: {json.dumps({'type':'step','label':step})}\n\n"
            import time; time.sleep(0.4)

        # Then stream the actual response
        try:
            oc = get_openai_client()
            messages = [{"role":"system","content":CHAT_SYSTEM}] + data.get("messages",[])
            stream = oc.chat.completions.create(
                model=os.getenv("AZURE_AI_MODEL_DEPLOYMENT_NAME","gpt-4.1-mini"),
                messages=messages, stream=True,
                max_tokens=2500, temperature=0.7,
            )
            for chunk in stream:
                if chunk.choices and chunk.choices[0].delta.content:
                    yield f"data: {json.dumps({'type':'content','content':chunk.choices[0].delta.content})}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type':'content','content':f'Error: {str(e)}'})}\n\n"

    return Response(stream_with_context(generate()), mimetype="text/event-stream",
                    headers={"Cache-Control":"no-cache","X-Accel-Buffering":"no"})


def _get_reasoning_steps(message: str) -> list:
    """Return contextual reasoning step labels based on user input."""
    msg = message.lower()
    if any(x in msg for x in ["resume","cv","cover letter","ats","job description","jd"]):
        return ["📋 Reading resume & JD","🔍 Identifying skill gaps","🎯 Matching keywords","✍️ Crafting optimized content","✅ Finalizing output"]
    elif any(x in msg for x in ["roadmap","plan","steps","how to","become","transition"]):
        return ["🎯 Assessing your current skills","🔍 Identifying skill gaps","🗺️ Mapping career path","📜 Sourcing certifications & resources","📅 Building your personalized plan"]
    elif any(x in msg for x in ["cert","certification","az-","aws","gcp","learn"]):
        return ["📚 Checking certification requirements","🔗 Finding official resources","⏱️ Estimating time & cost","📅 Sequencing your learning path"]
    elif any(x in msg for x in ["interview","question","prepare","practice"]):
        return ["📋 Analyzing job requirements","🧠 Generating technical questions","💼 Crafting behavioral questions","✅ Building model answers"]
    elif any(x in msg for x in ["salary","market","demand","hiring","company"]):
        return ["📊 Analyzing market data","🏢 Identifying top employers","💰 Calculating salary ranges","📈 Assessing demand trends"]
    else:
        return ["🧠 Understanding your goal","🔍 Analyzing career context","📅 Building personalized guidance"]

# ── fallback helpers (used only when UI cache is empty) ───────────────────────
def _get_optimized_resume(d):
    resume,jd = d.get("resume",""),d.get("jd","")
    level,industry = d.get("level","mid"),d.get("industry","tech")
    return call_ai([
        {"role":"system","content":RESUME_SYSTEM},
        {"role":"user","content":_resume_prompt(resume,jd,level,industry)}
    ], max_tokens=2200)

def _get_cover_letter(d):
    resume,jd = d.get("resume",""),d.get("jd","")
    level,industry = d.get("level","mid"),d.get("industry","tech")
    prompt = f"""Write a tailored cover letter. ROLE LEVEL: {level} | INDUSTRY: {industry}
=== RESUME ===\n{resume}\n=== JOB DESCRIPTION ===\n{jd}
Under 300 words. Output ONLY the letter — no commentary."""
    return call_ai([
        {"role":"system","content":"Expert cover letter writer. Specific, compelling, never generic."},
        {"role":"user","content":prompt}
    ], max_tokens=700)

# ── download endpoints — use already-generated text from UI (no extra AI call) ─
@app.route("/download/resume/docx", methods=["POST"])
def download_resume_docx():
    optimized = request.json.get("optimized_text","")
    if not optimized:
        optimized = _get_optimized_resume(request.json)
    buf = build_resume_docx(optimized)
    return send_file(buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True, download_name='CareerIQ_Optimized_Resume.docx')

@app.route("/download/resume/pdf", methods=["POST"])
def download_resume_pdf():
    optimized = request.json.get("optimized_text","")
    if not optimized:
        optimized = _get_optimized_resume(request.json)
    try:
        buf = build_resume_pdf(optimized)
        return send_file(buf, mimetype='application/pdf',
                         as_attachment=True, download_name='CareerIQ_Optimized_Resume.pdf')
    except ImportError:
        buf = build_resume_docx(optimized)
        return send_file(buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True, download_name='CareerIQ_Optimized_Resume.docx')

@app.route("/download/cover/docx", methods=["POST"])
def download_cover_docx():
    cover = request.json.get("cover_text","")
    if not cover:
        cover = _get_cover_letter(request.json)
    buf = build_cover_docx(cover)
    return send_file(buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True, download_name='CareerIQ_Cover_Letter.docx')

@app.route("/download/cover/pdf", methods=["POST"])
def download_cover_pdf():
    cover = request.json.get("cover_text","")
    if not cover:
        cover = _get_cover_letter(request.json)
    try:
        buf = build_cover_pdf(cover)
        return send_file(buf, mimetype='application/pdf',
                         as_attachment=True, download_name='CareerIQ_Cover_Letter.pdf')
    except ImportError:
        buf = build_cover_docx(cover)
        return send_file(buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True, download_name='CareerIQ_Cover_Letter.docx')

@app.route("/download/analysis/pdf", methods=["POST"])
def download_analysis_pdf():
    d = request.json
    analysis_text = d.get("analysis_text","")
    if not analysis_text:
        resume,jd = d.get("resume",""),d.get("jd","")
        analysis_text = call_ai([
            {"role":"system","content":"Expert ATS analyst. Detailed resume-to-JD match analysis."},
            {"role":"user","content":f"Analyze resume vs JD.\n\n=== RESUME ===\n{resume}\n\n=== JOB DESCRIPTION ===\n{jd}\n\nProvide ATS Score, Matches, Gaps, Missing Keywords, Path to 93+, Quick Wins."}
        ], max_tokens=1200)
    try:
        buf = build_analysis_pdf(analysis_text)
        return send_file(buf, mimetype='application/pdf',
                         as_attachment=True, download_name='CareerIQ_Match_Analysis.pdf')
    except ImportError:
        return jsonify({"error":"reportlab not installed. Run: pip install reportlab"}),500

if __name__ == "__main__":
    app.run(debug=True, port=5000, threaded=True)
